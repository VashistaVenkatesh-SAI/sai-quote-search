"""
SAI Module 1 BOM Generator
Upload quote PDFs or type specifications for instant BOM generation
Uses AI_Chatbot_Training_Module1_Assembly_Selection.docx for matching
"""
import streamlit as st
import openai
from azure.core.credentials import AzureKeyCredential
from azure.search.documents import SearchClient
from azure.search.documents.models import VectorizedQuery
import json
import re
import io

# Module 1 Matcher
try:
    from Module1Matcher import match_from_user_input, get_matcher, match_quote_to_assembly
    MODULE1_AVAILABLE = True
except ImportError:
    MODULE1_AVAILABLE = False

# PDF Processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Configuration
SEARCH_ENDPOINT = st.secrets["SEARCH_ENDPOINT"]
SEARCH_KEY = st.secrets["SEARCH_KEY"]
INDEX_NAME = st.secrets["INDEX_NAME"]
AZURE_OPENAI_ENDPOINT = st.secrets["AZURE_OPENAI_ENDPOINT"]
AZURE_OPENAI_KEY = st.secrets["AZURE_OPENAI_KEY"]
AZURE_OPENAI_DEPLOYMENT = st.secrets["AZURE_OPENAI_DEPLOYMENT"]
AZURE_OPENAI_DEPLOYMENT_EMBEDDINGS = st.secrets["AZURE_OPENAI_DEPLOYMENT_EMBEDDINGS"]
AUTHORIZED_USERS = st.secrets["AUTHORIZED_USERS"]

openai.api_type = "azure"
openai.api_key = AZURE_OPENAI_KEY
openai.api_base = AZURE_OPENAI_ENDPOINT
openai.api_version = "2024-02-01"

# Page config
st.set_page_config(
    page_title="SAI Module 1 BOM Generator",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS Styling
st.markdown("""
<style>
    :root {
        --bg-primary: #1e1e1e;
        --bg-secondary: #2d2d2d;
        --text-primary: #e8e8e8;
        --text-secondary: #b0b0b0;
        --accent-blue: #2563EB;
        --accent-green: #10B981;
        --border-color: #404040;
    }
    
    .stApp {
        background-color: #1e1e1e;
        color: #e8e8e8;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    .main-header {
        background: #2d2d2d;
        padding: 1.5rem 2rem;
        border-radius: 12px;
        margin-bottom: 2rem;
        border: 1px solid #404040;
    }
    
    .main-header h1 {
        color: #e8e8e8;
        margin: 0;
        font-size: 1.75rem;
        font-weight: 600;
    }
    
    .main-header p {
        color: #b0b0b0;
        margin: 0.5rem 0 0 0;
        font-size: 0.95rem;
    }
    
    .module1-badge {
        background: #10B981;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 12px;
        font-size: 0.75rem;
        font-weight: 600;
        display: inline-block;
        margin-left: 1rem;
    }
    
    .user-message {
        background: #2563EB;
        color: white;
        padding: 1rem 1.25rem;
        border-radius: 18px;
        margin: 1.5rem 0;
        max-width: 85%;
        margin-left: auto;
        font-size: 0.95rem;
        line-height: 1.5;
    }
    
    .assistant-message {
        background: #2d2d2d;
        color: #e8e8e8;
        padding: 1.25rem 1.5rem;
        border-radius: 18px;
        margin: 1.5rem 0;
        max-width: 85%;
        border: 1px solid #404040;
        font-size: 0.95rem;
        line-height: 1.6;
    }
    
    .bom-card {
        background: linear-gradient(135deg, #2d2d2d 0%, #3a3a3a 100%);
        border-radius: 16px;
        padding: 2rem;
        margin: 1.5rem 0;
        border: 2px solid #10B981;
        box-shadow: 0 8px 24px rgba(16, 185, 129, 0.2);
    }
    
    .bom-header {
        display: flex;
        align-items: center;
        gap: 1rem;
        margin-bottom: 1.5rem;
        padding-bottom: 1rem;
        border-bottom: 1px solid #404040;
    }
    
    .bom-title {
        font-size: 1.5rem;
        font-weight: 700;
        color: #10B981;
    }
    
    .component-item {
        background: #2d2d2d;
        padding: 0.875rem 1.25rem;
        border-radius: 8px;
        margin: 0.5rem 0;
        border-left: 3px solid #10B981;
    }
    
    .component-number {
        font-family: 'Courier New', monospace;
        color: #10B981;
        font-weight: 600;
    }
    
    .status-exact {
        background: #10B981;
        color: white;
        padding: 0.375rem 0.875rem;
        border-radius: 16px;
        font-size: 0.8rem;
        font-weight: 600;
    }
    
    .stButton > button {
        background: #2563EB;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 500;
        transition: all 0.2s;
    }
    
    .stButton > button:hover {
        background: #1d4ed8;
        box-shadow: 0 4px 12px rgba(37, 99, 235, 0.3);
    }
    
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #2d2d2d;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #404040;
        border-radius: 4px;
    }
</style>
""", unsafe_allow_html=True)

@st.cache_data
def load_training_document():
    """Load training document from GitHub repo"""
    try:
        from docx import Document
        doc = Document('AI_Chatbot_Training_Module1_Assembly_Selection.docx')
        
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        # Also get tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    full_text.append(' | '.join(cells))
        
        content = '\n'.join(full_text)
        return content
        
    except Exception as e:
        return None

def check_password():
    """Authentication"""
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
    
    if st.session_state.authenticated:
        return True
    
    st.markdown("""
    <div class="main-header">
        <h1>SAI Module 1 BOM Generator</h1>
        <p>Sign in to access the automated BOM generation system</p>
    </div>
    """, unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("### Welcome back")
        username = st.text_input("Username", key="username")
        password = st.text_input("Password", type="password", key="password")
        
        if st.button("Sign in", use_container_width=True):
            if username in AUTHORIZED_USERS and AUTHORIZED_USERS[username] == password:
                st.session_state.authenticated = True
                st.session_state.current_user = username
                st.success("Signed in successfully")
                st.rerun()
            else:
                st.error("Invalid credentials")
        
        st.markdown("---")
        st.caption("For authorized SAI personnel only")
    
    return False

# Initialize session state
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'search_client' not in st.session_state:
    st.session_state.search_client = SearchClient(
        endpoint=SEARCH_ENDPOINT,
        index_name=INDEX_NAME,
        credential=AzureKeyCredential(SEARCH_KEY)
    )

# Check authentication
if not check_password():
    st.stop()

def extract_text_from_pdf(pdf_file):
    """Extract text from PDF"""
    try:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def extract_specs_from_text(text):
    """Extract specifications and explain matching using training document"""
    
    # Load ACTUAL training document
    training_doc = load_training_document()
    
    if training_doc:
        training_content = training_doc[:15000]  # Use more of the doc
    else:
        training_content = "Training document not available."

    system_prompt = f"""You are an expert at matching switchgear quotes to Module 1 assemblies using the training document.

COMPLETE TRAINING DOCUMENT:
{training_content}

YOUR TASK:
1. Read the quote and extract specifications
2. Compare to the examples and patterns in the training document
3. Determine which assembly from the training doc matches
4. EXPLAIN your reasoning by referencing the training document

When you find a match, explain it like this:
"Based on the training document, this quote matches Assembly [NUMBER] because:
- The training doc Example [X] shows that quotes with [these features] match Assembly [NUMBER]
- Your quote has: [extracted features]
- According to the training document's matching rules, [explain why it matches]"

Extract specs as JSON AND provide explanation:
{{
  "sections": [
    {{
      "identifier": "Section 101",
      "dimensions": {{"height": "90", "width": "40", "depth": "60"}},
      "main_circuit_breaker": {{"type": "ABB SACE Emax 6.2", "quantity": 1}}
    }}
  ],
  "special_construction_requirements": ["fixed mount", "front and rear access"],
  "reasoning": "Based on training document Example 1 (page X), quotes with 90H x 40W x 60D and Emax 6.2 match Assembly 123456-0100-101. The training doc shows this configuration requires fixed mount and front/rear access, which this quote has."
}}

CRITICAL: Always reference the training document in your reasoning. Show which example or pattern you used."""

    user_prompt = f"""Using the training document as your guide, extract specs from this quote AND explain which assembly it matches and WHY based on the training document.

Quote:
{text[:12000]}

Return JSON with specs AND reasoning that references the training document."""
    
    try:
        response = openai.ChatCompletion.create(
            engine=AZURE_OPENAI_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.0,
            max_tokens=2500,
            timeout=30
        )
        
        ai_response = response.choices[0].message.content.strip()
        
        # Clean JSON
        ai_response = ai_response.replace("```json", "").replace("```", "").strip()
        start = ai_response.find('{')
        end = ai_response.rfind('}') + 1
        if start != -1 and end > start:
            ai_response = ai_response[start:end]
        
        # Remove trailing commas
        for _ in range(5):
            ai_response = re.sub(r',(\s*[}\]])', r'\1', ai_response)
        
        parsed = json.loads(ai_response)
        return parsed
        
    except Exception as e:
        st.error(f"Error extracting specs: {e}")
        return None

def display_bom_card(bom_data):
    """Display Module 1 BOM card with match explanations"""
    
    status = bom_data.get('status')
    matched_assemblies = bom_data.get('matched_assemblies', [])
    extracted_features = bom_data.get('extracted_features', {})
    
    # Show selection buttons if multiple/no exact match
    if status in ['ambiguous', 'no_match'] and matched_assemblies:
        # Show detected specs
        st.markdown("### 📋 Detected Specifications:")
        det_col1, det_col2, det_col3, det_col4 = st.columns(4)
        with det_col1:
            st.metric("Height", f"{extracted_features.get('height', '?')}\"")
        with det_col2:
            st.metric("Width", f"{extracted_features.get('width', '?')}\"")
        with det_col3:
            st.metric("Depth", f"{extracted_features.get('depth', '?')}\"")
        with det_col4:
            breaker = extracted_features.get('breaker_type', 'Not specified')
            st.metric("Breaker", breaker[:20] if len(breaker) > 20 else breaker)
        
        st.markdown("---")
        st.markdown("### 🔍 Select an assembly:")
        
        matcher = get_matcher()
        num_assemblies = min(len(matched_assemblies), 9)
        cols_per_row = 3
        
        for i in range(0, num_assemblies, cols_per_row):
            cols = st.columns(cols_per_row)
            for j in range(cols_per_row):
                idx = i + j
                if idx < num_assemblies:
                    assembly_num = matched_assemblies[idx]
                    specs = matcher.assembly_specs[assembly_num]
                    
                    with cols[j]:
                        # Calculate match score
                        match_score = 0
                        total_possible = 4
                        
                        height_match = extracted_features.get('height') == specs['height']
                        width_match = extracted_features.get('width') == specs['width']
                        depth_match = extracted_features.get('depth') == specs['depth']
                        breaker_match = False
                        
                        if extracted_features.get('breaker_type'):
                            breaker_match = extracted_features['breaker_type'].upper() in specs['breaker_type'].upper() or specs['breaker_type'].upper() in extracted_features['breaker_type'].upper()
                        
                        if height_match:
                            match_score += 1
                        if width_match:
                            match_score += 1
                        if depth_match:
                            match_score += 1
                        if breaker_match:
                            match_score += 1
                        
                        match_pct = int((match_score / total_possible) * 100)
                        
                        # Button
                        if st.button(f"**{assembly_num}**\n{match_pct}% Match", key=f"select_{assembly_num}_{idx}", use_container_width=True):
                            selected_bom = matcher.generate_bom(assembly_num)
                            
                            st.session_state.messages.append({
                                "role": "user",
                                "content": f"Show BOM for {assembly_num}"
                            })
                            
                            st.session_state.messages.append({
                                "role": "assistant",
                                "content": f"✅ Showing BOM for Assembly {assembly_num}",
                                "module1_result": {
                                    'status': 'exact_match',
                                    'bom': selected_bom,
                                    'message': f"✅ BOM for {assembly_num}"
                                },
                                "type": "module1"
                            })
                            
                            st.rerun()
                        
                        # Match details
                        st.caption("**Match Details:**")
                        st.caption(f"{'✅' if height_match else '❌'} Height: {specs['height']}\"")
                        st.caption(f"{'✅' if width_match else '❌'} Width: {specs['width']}\"")
                        st.caption(f"{'✅' if depth_match else '❌'} Depth: {specs['depth']}\"")
                        st.caption(f"{'✅' if breaker_match else '❌'} Breaker: {specs['breaker_type'][:20]}...")
                        st.caption(f"Mount: {specs['mount']}")
                        st.caption(f"Access: {specs['access']}")
        
        return  # Exit early
    
    # Show BOM card for exact match
    if not bom_data or 'bom' not in bom_data or not bom_data['bom']:
        return
    
    bom = bom_data['bom']
    badge_html = '<span class="status-exact">✅ Match Found</span>'
    
    st.markdown(f"""
    <div class="bom-card">
        <div class="bom-header">
            <div class="bom-title">🔧 Module 1 BOM</div>
            {badge_html}
        </div>
        <div style="font-size: 1.125rem; color: #e8e8e8; font-weight: 600;">
            Assembly: {bom['assembly_number']}
        </div>
        <div style="color: #b0b0b0; margin-top: 0.5rem;">{bom['project']}</div>
    </div>
    """, unsafe_allow_html=True)
    
    # Specifications
    specs = bom['specifications']
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("**Dimensions**")
        st.code(f"{specs['height']}\"H x {specs['width']}\"W x {specs['depth']}\"D")
    
    with col2:
        st.markdown("**Breaker**")
        st.code(f"{specs['breaker_type']}")
    
    with col3:
        st.markdown("**Mount**")
        st.code(specs['mount'])
    
    with col4:
        st.markdown("**Access**")
        st.code(specs['access'])
    
    st.markdown(f"**Total Components:** `{bom['total_parts']} parts`")
    
    # Components list
    with st.expander(f"📋 View all {bom['total_parts']} components", expanded=False):
        for i, comp in enumerate(bom['components'], 1):
            st.markdown(f"""
            <div class="component-item">
                <span style="color: #b0b0b0; font-size: 0.85rem;">#{i}</span>
                <span class="component-number">{comp['part_number']}</span>
                <span style="color: #b0b0b0; margin-left: 1rem;">Qty: {comp['quantity']}</span>
                <div style="color: #b0b0b0; font-size: 0.85rem; margin-top: 0.25rem;">{comp.get('description', '')[:80]}</div>
            </div>
            """, unsafe_allow_html=True)
    
    # Export to CSV
    if st.button("📥 Export BOM to CSV", key=f"export_{bom['assembly_number']}"):
        csv_buffer = io.StringIO()
        csv_buffer.write("Item,Part Number,Description,Quantity\n")
        
        for i, comp in enumerate(bom['components'], 1):
            part_num = comp['part_number'].replace(',', ';')
            desc = comp.get('description', '').replace(',', ';').replace('\n', ' ')
            qty = comp['quantity']
            csv_buffer.write(f"{i},{part_num},{desc},{qty}\n")
        
        csv_str = csv_buffer.getvalue()
        
        st.download_button(
            label="📥 Download CSV",
            data=csv_str,
            file_name=f"{bom['assembly_number']}_BOM.csv",
            mime="text/csv",
            key=f"download_{bom['assembly_number']}"
        )

# Header
module1_badge = '<span class="module1-badge">BOM Generator</span>' if MODULE1_AVAILABLE else ''
st.markdown(f"""
<div class="main-header">
    <h1>SAI Module 1 {module1_badge}</h1>
    <p>Upload quote PDFs or type specifications for instant BOM generation</p>
</div>
""", unsafe_allow_html=True)

# Chat input
user_input = st.chat_input("Type specs for Module 1 BOM... (e.g., '90H x 40W x 60D, Emax 6.2')")

if user_input:
    st.session_state.messages.append({"role": "user", "content": user_input})
    
    if MODULE1_AVAILABLE:
        with st.spinner("Matching to Module 1 assembly..."):
            module1_result = match_from_user_input(user_input)
            
            st.session_state.messages.append({
                "role": "assistant",
                "content": module1_result['message'],
                "module1_result": module1_result,
                "type": "module1"
            })
    else:
        st.session_state.messages.append({
            "role": "assistant",
            "content": "Module 1 matching not available.",
            "type": "error"
        })

# Display chat history
for message in st.session_state.messages:
    if message["role"] == "user":
        st.markdown(f'<div class="user-message">{message["content"]}</div>', unsafe_allow_html=True)
    else:
        st.markdown(f'<div class="assistant-message">{message["content"]}</div>', unsafe_allow_html=True)
        
        if message.get("type") == "module1" and "module1_result" in message:
            display_bom_card(message["module1_result"])

# Sidebar
with st.sidebar:
    st.markdown(f"### 👤 {st.session_state.get('current_user', 'User')}")
    
    if st.button("🚪 Sign out"):
        st.session_state.authenticated = False
        st.session_state.current_user = None
        st.rerun()
    
    st.markdown("---")
    
    # PDF Upload
    if MODULE1_AVAILABLE and PDF_AVAILABLE:
        st.markdown("### 📤 Upload Quote PDF")
        st.caption("Get instant Module 1 BOM")
        
        uploaded_pdf = st.file_uploader(
            "Drop quote PDF here",
            type=['pdf'],
            help="Upload quote to auto-generate BOM",
            key="pdf_uploader"
        )
        
        if uploaded_pdf is not None:
            if st.button("🔧 Generate BOM", use_container_width=True):
                with st.spinner("📄 Reading PDF..."):
                    text = extract_text_from_pdf(uploaded_pdf)
                    
                    if text:
                        with st.spinner("🤖 Analyzing with training document..."):
                            specs_json = extract_specs_from_text(text)
                            
                            if specs_json:
                                # Show AI reasoning if available
                                if 'reasoning' in specs_json:
                                    st.info(f"**AI Analysis:**\n\n{specs_json['reasoning']}")
                                
                                with st.spinner("🔍 Matching to Module 1..."):
                                    module1_result = match_quote_to_assembly(specs_json)
                                    
                                    features = module1_result.get('extracted_features', {})
                                    feature_text = f"Detected: {features.get('height', '?')}\"H x {features.get('width', '?')}\"W x {features.get('depth', '?')}\"D"
                                    if features.get('breaker_type'):
                                        feature_text += f", {features['breaker_type']}"
                                    
                                    # Add AI reasoning to message if available
                                    ai_reasoning = specs_json.get('reasoning', '')
                                    full_message = module1_result['message']
                                    if ai_reasoning:
                                        full_message = f"{ai_reasoning}\n\n{full_message}"
                                    
                                    st.session_state.messages.append({
                                        "role": "user",
                                        "content": f"📄 {uploaded_pdf.name}\n{feature_text}"
                                    })
                                    
                                    st.session_state.messages.append({
                                        "role": "assistant",
                                        "content": full_message,
                                        "module1_result": module1_result,
                                        "type": "module1"
                                    })
                                    
                                    st.rerun()
                            else:
                                st.error("❌ Could not extract specs")
                    else:
                        st.error("❌ Could not read PDF")
        
        st.markdown("---")
    
    # Examples
    if MODULE1_AVAILABLE:
        st.markdown("### 💡 Example Queries")
        
        with st.expander("Try these"):
            if st.button("90H x 40W x 60D, Emax 6.2", use_container_width=True):
                st.session_state.messages.append({
                    "role": "user",
                    "content": "90H x 40W x 60D, Emax 6.2, fixed, front and rear"
                })
                st.rerun()
            
            if st.button("78H x 42W x 33D, Square D", use_container_width=True):
                st.session_state.messages.append({
                    "role": "user",
                    "content": "78H x 42W x 33D, Square D"
                })
                st.rerun()
        
        if st.button("📋 List All Assemblies", use_container_width=True):
            matcher = get_matcher()
            st.markdown("#### Available Assemblies:")
            for asm_num in matcher.assembly_specs.keys():
                specs = matcher.assembly_specs[asm_num]
                st.caption(f"**{asm_num}**")
                st.caption(f"{specs['height']}\"H x {specs['width']}\"W x {specs['depth']}\"D")
                st.markdown("")
        
        st.markdown("---")
    
    st.markdown("### 📊 Stats")
    st.caption(f"{len(st.session_state.messages)} messages")
    
    if st.button("🗑️ Clear chat"):
        st.session_state.messages = []
        st.rerun()
    
    st.markdown("---")
    st.caption("SAI Advanced Power Solutions")
    st.caption("Module 1 BOM Generator v3.0")
